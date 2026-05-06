import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _ts_to_seconds(ts: str) -> int:
    """Parse 'HH:MM:SS' or 'MM:SS' to whole seconds. Returns 0 on bad input."""
    if not ts:
        return 0
    parts = [p for p in str(ts).strip().split(':') if p.strip()]
    try:
        nums = [int(float(p)) for p in parts]
    except ValueError:
        return 0
    if len(nums) == 3:
        h, m, s = nums
    elif len(nums) == 2:
        h, m, s = 0, nums[0], nums[1]
    elif len(nums) == 1:
        h, m, s = 0, 0, nums[0]
    else:
        return 0
    return h * 3600 + m * 60 + s


def _yt_link(yt_id: str, ts: str) -> str:
    """Build a deep-link to a YouTube video at the given timestamp.
    Returns '' if we don't have a video id."""
    if not yt_id:
        return ''
    secs = _ts_to_seconds(ts)
    return f"https://youtu.be/{yt_id}?t={secs}s"


def _add_hyperlink(paragraph, url: str, text: str,
                   color: str = '1A56B0', size: int = 8, bold: bool = True):
    """Add a clickable hyperlink run inside a python-docx paragraph.

    python-docx has no first-class API for this, so we build the underlying
    XML directly: relate_to() registers the URL in the document's
    relationship table, then we wrap a styled <w:r> in a <w:hyperlink>.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), color); rPr.append(color_el)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size * 2)); rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# Colours
C_GREY     = RGBColor(0x80, 0x80, 0x80)
C_DARK     = RGBColor(0x33, 0x33, 0x33)
C_BLUE     = RGBColor(0x1a, 0x56, 0xb0)
C_PURPLE   = RGBColor(0x6a, 0x0d, 0xad)
C_GREEN    = RGBColor(0x1a, 0x7a, 0x3c)

BG_HEADER  = 'FFD0D0'   # pink  – column headers
BG_SECTION = 'E8EAF6'   # indigo tint – section row
BG_VO      = 'FFF9E6'   # warm cream – VO rows
BG_VISUAL  = 'E8F5E9'   # light green – Visual rows


def export_to_docx(data: dict, output_dir: str, title: str, yt_id: str = '') -> str:
    doc = Document()

    # A4 landscape (297 × 210 mm) — standard non-US paper size
    sec = doc.sections[0]
    sec.page_width   = Mm(297)
    sec.page_height  = Mm(210)
    sec.left_margin  = sec.right_margin = Mm(15)
    sec.top_margin   = sec.bottom_margin = Mm(15)

    # ── Title block ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get('title', title))
    r.bold = True; r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Duration: {data.get('total_duration', 'N/A')}")
    r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = C_GREY

    doc.add_paragraph()

    # ── Main two-column table ─────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # A4 landscape minus margins ≈ 267mm; split evenly between VO and Visuals.
    _set_col_width(tbl, 0, Mm(133))
    _set_col_width(tbl, 1, Mm(133))

    # Column header row
    hdr = tbl.rows[0].cells
    _bg(hdr[0], BG_HEADER); _bg(hdr[1], BG_HEADER)
    hdr[0].paragraphs[0].add_run('VO').bold = True
    hdr[1].paragraphs[0].add_run('Visuals  (plays AFTER the VO)').bold = True

    for section in data.get('sections', []):
        # Section title spanning both columns
        sr = tbl.add_row()
        mc = sr.cells[0].merge(sr.cells[1])
        _bg(mc, BG_SECTION)
        run = mc.paragraphs[0].add_run(section.get('title', '').upper())
        run.bold = True; run.font.size = Pt(10); run.font.color.rgb = C_BLUE

        for beat in section.get('beats', []):
            vo   = beat.get('vo', {})
            viz  = beat.get('visual') or beat.get('visual_after') or {}

            row = tbl.add_row()
            vo_cell, vis_cell = row.cells[0], row.cells[1]

            # ── VO cell ──────────────────────────────────────────────────────
            _bg(vo_cell, BG_VO)
            vp = vo_cell.paragraphs[0]

            # Timestamp — hyperlink when we have a YouTube id, plain otherwise.
            vo_ts_s = vo.get('timestamp_start', '')
            vo_ts_e = vo.get('timestamp_end', '')
            ts_text = f"[{vo_ts_s} – {vo_ts_e}]"
            link = _yt_link(yt_id, vo_ts_s)
            if link:
                _add_hyperlink(vp, link, ts_text, color='1A56B0', size=8, bold=True)
                vp.add_run('\n')
            else:
                r = vp.add_run(ts_text + '\n')
                r.font.size = Pt(8); r.font.color.rgb = C_GREY

            # Tone
            tone = vo.get('tone', '')
            if tone:
                r = vp.add_run(tone + '\n'); r.italic = True
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66,0x66,0x66)

            # VO text
            vp.add_run(vo.get('text', ''))

            # ── Visual cell ──────────────────────────────────────────────────
            _bg(vis_cell, BG_VISUAL)
            vizp = vis_cell.paragraphs[0]

            # Timestamp — hyperlink when we have a YouTube id.
            vis_ts_s = viz.get('timestamp_start', '')
            vis_ts_e = viz.get('timestamp_end', '')
            vts = f"({vis_ts_s} – {vis_ts_e})"
            link = _yt_link(yt_id, vis_ts_s)
            if link:
                _add_hyperlink(vizp, link, vts, color='15803D', size=9, bold=True)
                vizp.add_run('\n')
            else:
                r = vizp.add_run(vts + '\n'); r.bold = True
                r.font.size = Pt(9); r.font.color.rgb = C_PURPLE

            # Description
            desc = viz.get('description', '')
            if desc:
                r = vizp.add_run(desc + '\n\n')
                r.font.size = Pt(10)

            # On-screen text
            ost = viz.get('on_screen_text', '')
            if ost and ost.upper() != 'NONE':
                r = vizp.add_run('On-screen text: '); r.bold = True; r.font.size = Pt(9)
                r = vizp.add_run(ost + '\n'); r.font.size = Pt(9); r.font.color.rgb = C_DARK

            # Audio notes
            audio = viz.get('audio_notes', '')
            if audio:
                r = vizp.add_run('Audio: '); r.bold = True; r.font.size = Pt(9)
                r = vizp.add_run(audio + '\n'); r.font.size = Pt(9); r.font.color.rgb = C_GREY

            # Visual summary
            vsumm = viz.get('summary', '')
            if vsumm:
                r = vizp.add_run('\nSummary: '); r.bold = True; r.italic = True; r.font.size = Pt(9)
                r = vizp.add_run(vsumm); r.italic = True; r.font.size = Pt(9)
                r.font.color.rgb = C_GREEN

    # ── Summary page ──────────────────────────────────────────────────────────
    doc.add_page_break()

    h = doc.add_heading('Summary & Highlights', 1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _labeled(doc, 'Full Summary', data.get('summary', ''))

    peaks = data.get('peak_moments', [])
    if peaks:
        doc.add_paragraph()
        _labeled(doc, 'Peak Moments', '')
        for pm in peaks:
            p = doc.add_paragraph(style='List Bullet')
            ts = pm.get('timestamp', '')
            link = _yt_link(yt_id, ts)
            if link:
                _add_hyperlink(p, link, f"[{ts}]", color='1A56B0', size=10, bold=True)
                p.add_run(' ')
            else:
                p.add_run(f"[{ts}] ").bold = True
            p.add_run(pm.get('description', ''))

    highlights = data.get('highlights', [])
    if highlights:
        doc.add_paragraph()
        _labeled(doc, 'Key Highlights', '')
        for hl in highlights:
            doc.add_paragraph(hl, style='List Bullet')

    out = os.path.join(output_dir, f"{_safe(title)}_analysis.docx")
    doc.save(out)
    print(f"DOCX saved → {out}")
    return out


def export_to_txt(data: dict, output_dir: str, title: str, yt_id: str = '') -> str:
    # 80 chars wide — fits A4 portrait at 12pt monospace with 15mm margins.
    W = 80
    sep  = '=' * W
    thin = '─' * W

    def ts(ts_str):
        """Render a timestamp; append a clickable YT URL when we have an id."""
        link = _yt_link(yt_id, ts_str)
        return f"{ts_str}  →  {link}" if link else ts_str

    lines = [sep, f"VIDEO ANALYSIS: {data.get('title', title)}",
             f"Duration: {data.get('total_duration','N/A')}"]
    if yt_id:
        lines.append(f"Source:   https://youtu.be/{yt_id}")
    lines += [sep, '']

    for section in data.get('sections', []):
        lines += ['', thin, f"  SECTION: {section.get('title','').upper()}", thin, '']

        for i, beat in enumerate(section.get('beats', []), 1):
            vo  = beat.get('vo', {})
            viz = beat.get('visual') or beat.get('visual_after') or {}

            vo_ts = vo.get('timestamp_start','')
            vi_ts = viz.get('timestamp_start','')

            lines.append(f"  BEAT {i}")
            lines.append(f"  ┌─ VO  [{vo_ts} – {vo.get('timestamp_end','')}]")
            if yt_id and vo_ts:
                lines.append(f"  │   Watch  : {_yt_link(yt_id, vo_ts)}")
            if vo.get('tone'):
                lines.append(f"  │   Tone   : {vo['tone']}")
            lines.append(f"  │   Text   : {vo.get('text','')}")
            lines.append( "  │")
            lines.append(f"  └─ VISUAL  ({vi_ts} – {viz.get('timestamp_end','')})")
            if yt_id and vi_ts:
                lines.append(f"      Watch       : {_yt_link(yt_id, vi_ts)}")
            lines.append(f"      Description : {viz.get('description','')}")
            if viz.get('on_screen_text') and viz['on_screen_text'].upper() != 'NONE':
                lines.append(f"      On-screen   : {viz['on_screen_text']}")
            if viz.get('audio_notes'):
                lines.append(f"      Audio       : {viz['audio_notes']}")
            if viz.get('summary'):
                lines.append(f"      Summary     : {viz['summary']}")
            lines.append('')

    lines += ['', sep, 'SUMMARY', sep, data.get('summary',''), '']

    if data.get('peak_moments'):
        lines += [thin, 'PEAK MOMENTS', thin]
        for pm in data['peak_moments']:
            pts = pm.get('timestamp','')
            lines.append(f"  [{ts(pts)}]  {pm.get('description','')}")
        lines.append('')

    if data.get('highlights'):
        lines += [thin, 'KEY HIGHLIGHTS', thin]
        for hl in data['highlights']:
            lines.append(f"  • {hl}")

    out = os.path.join(output_dir, f"{_safe(title)}_analysis.txt")
    with open(out, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"TXT saved  → {out}")
    return out


# ── helpers ──────────────────────────────────────────────────────────────────

def _bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def _set_col_width(table, idx: int, width):
    for row in table.rows:
        row.cells[idx].width = width


def _labeled(doc, label: str, content: str):
    p = doc.add_paragraph()
    r = p.add_run(label + '\n'); r.bold = True; r.font.size = Pt(12)
    if content:
        p.add_run(content)


def _safe(name: str) -> str:
    return re.sub(r'[^\w\s\-]', '', name).strip().replace(' ', '_')[:60]
